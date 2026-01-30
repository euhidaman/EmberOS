"""
Document Reader Tool for EmberOS.

Reads and extracts content from various document formats:
- Plain text (.txt)
- Markdown (.md)
- PDF (.pdf)
- Word documents (.docx)
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any

from emberos.tools.base import (
    BaseTool, ToolResult, ToolManifest, ToolParameter,
    ToolCategory, RiskLevel
)
from emberos.tools.registry import register_tool

logger = logging.getLogger(__name__)


@register_tool
class DocumentReaderTool(BaseTool):
    """Tool for reading and extracting content from various document formats."""

    @property
    def manifest(self) -> ToolManifest:
        return ToolManifest(
            name="document.read",
            description="Read and extract content from documents (TXT, MD, PDF, DOCX)",
            category=ToolCategory.FILESYSTEM,
            risk_level=RiskLevel.LOW,
            parameters=[
                ToolParameter(
                    name="filepath",
                    description="Path to the document file",
                    type="string",
                    required=True
                )
            ]
        )

    def __init__(self):
        """Initialize document reader."""
        super().__init__()
        self.max_size = 10 * 1024 * 1024  # 10MB limit

    async def execute(self, filepath: str) -> ToolResult:
        """
        Read a document and extract its content.

        Args:
            filepath: Path to the document

        Returns:
            ToolResult with document content and metadata
        """
        try:
            # Expand path
            filepath = os.path.expanduser(filepath)

            # Check if file exists
            if not os.path.exists(filepath):
                return ToolResult(
                    success=False,
                    result=None,
                    error=f"File not found: {filepath}"
                )

            # Check file size
            file_size = os.path.getsize(filepath)
            if file_size > self.max_size:
                return ToolResult(
                    success=False,
                    result=None,
                    error=f"File too large: {file_size / (1024*1024):.1f}MB (max: {self.max_size / (1024*1024):.1f}MB)"
                )

            # Get file extension
            ext = Path(filepath).suffix.lower()

            # Get basic metadata
            metadata = {
                "filename": os.path.basename(filepath),
                "size": file_size,
                "extension": ext,
                "path": filepath
            }

            # Read based on file type
            if ext in ['.txt', '.md', '.markdown']:
                content = self._read_text_file(filepath)
            elif ext == '.pdf':
                content = self._read_pdf(filepath)
            elif ext in ['.docx', '.doc']:
                content = self._read_docx(filepath)
            else:
                return ToolResult(
                    success=False,
                    result={"metadata": metadata},
                    error=f"Unsupported file format: {ext}. Supported: .txt, .md, .pdf, .docx"
                )

            if content is None:
                return ToolResult(
                    success=False,
                    result={"metadata": metadata},
                    error="Failed to extract content from file"
                )

            metadata["length"] = len(content)
            metadata["word_count"] = len(content.split())

            return ToolResult(
                success=True,
                result={
                    "success": True,
                    "content": content,
                    "metadata": metadata
                },
                error=None
            )

        except Exception as e:
            logger.exception(f"Error reading document {filepath}: {e}")
            return ToolResult(
                success=False,
                result=None,
                error=str(e)
            )

    def _read_text_file(self, filepath: str) -> Optional[str]:
        """Read plain text or markdown file."""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encodings
            try:
                with open(filepath, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Failed to read text file {filepath}: {e}")
                return None
        except Exception as e:
            logger.error(f"Failed to read text file {filepath}: {e}")
            return None

    def _read_pdf(self, filepath: str) -> Optional[str]:
        """Read PDF file and extract text."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(filepath)
            text_parts = []

            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue

            if not text_parts:
                return None

            return "\n\n".join(text_parts)

        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            return None
        except Exception as e:
            logger.error(f"Failed to read PDF {filepath}: {e}")
            return None

    def _read_docx(self, filepath: str) -> Optional[str]:
        """Read Word document and extract text."""
        try:
            from docx import Document

            doc = Document(filepath)
            text_parts = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            if not text_parts:
                return None

            return "\n\n".join(text_parts)

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"Failed to read DOCX {filepath}: {e}")
            return None

